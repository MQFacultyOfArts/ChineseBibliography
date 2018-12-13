#!/usr/bin/env python3
# By Brian Ballsun-Stanton
# GPL v3 

import google
from collections import defaultdict
from pprint import pprint
from docx import Document
from docx.shared import Cm
import os
import shutil

sheet = google.getBibliographicItems()

shutil.rmtree('output', ignore_errors=True)
os.makedirs('output')

output = {}

for item in sheet:
	pprint(item)

	if ', ' in item.get('Library'):
		originalLibrary = item.get('Library')
		for library in originalLibrary.split(', '):
			if library not in output:
				output[library] = {}
			if item.get('CategoryName', '') not in output[library]:
				output[library][item.get('CategoryName', '')] = []	
			output[library][item.get('CategoryName', '')].append(item)


	else:
		if item.get('Library', '') not in output:
			output[item.get('Library', '')] = {}
		if item.get('CategoryName', '') not in output[item.get('Library', '')]:
			output[item.get('Library', '')][item.get('CategoryName', '')] = []	
		output[item.get('Library', '')][item.get('CategoryName', '')].append(item)

pprint(output)
for library in output:
	document = Document()
	for category in output[library]:
		document.add_heading(category, 1)
		for item in output[library][category]:
			p = document.add_paragraph()
			p.add_run("{}{}".format(item.get('Book_Title', ''), item.get('Volume_Count', ''))).bold = True
			p.add_run(" {}".format(item.get('Accession_Number', '')))

			p2 = document.add_paragraph("〔{}〕{}".format(item.get('Author_1_Period', '?'), item.get('Authorship_String_1', '')))
			paragraph_format = p2.paragraph_format
			paragraph_format.left_indent = Cm(0.5)

			if item.get('Book_title_relative_to_author_1', ''):
				p2.add_run("《{}》".format(item.get('Book_title_relative_to_author_1', '')))

			if item.get('Authorship_String_2', ''):
				p2.add_run("〔{}〕{}".format(item.get('Author_2_Period', ''), item.get('Authorship_String_2', '')))

			p3 = document.add_paragraph("{}".format(item.get('Date_of_Publication', '')))
			paragraph_format = p3.paragraph_format
			paragraph_format.left_indent = Cm(0.5)

			if 'Western_Year' in item and item.get('Western_Year', ''):
				p3.add_run("（{}）".format(item.get('Western_Year', '')))

			if 'Name_of_Publisher' in item and item.get('Name_of_Publisher', ''):
				p3.add_run(" {}".format(item.get('Name_of_Publisher', '')))

			if 'Note' in item:
				p4 = document.add_paragraph("{}".format(item.get('Note', '')))
				paragraph_format = p4.paragraph_format
				paragraph_format.left_indent = Cm(0.5)


	document.save('output/{}.docx'.format(library))

"""

Repeating:
	[Bold:]Book title Volume Count " " [Not bold:]Accession Number
	(Period) Commentator's Name Etal Ownership verb 《Original Book Title》(Period) Commentator's Name Etal Ownership verb
	Date of Publication (Western Year) Name of Publisher Edition
	Note
	"""

